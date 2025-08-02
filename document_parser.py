import os
import re
from typing import List, Dict, Optional
from urllib.parse import urlparse

# Use PyMuPDF (fitz) as primary PDF parser
try:
    import fitz  # PyMuPDF
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False
    print("Warning: PyMuPDF not available")

# Fallback PDF parser
try:
    import PyPDF2
    PYPDF2_AVAILABLE = True
except ImportError:
    PYPDF2_AVAILABLE = False
    print("Warning: PyPDF2 not available")

# DOCX parser
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("Warning: python-docx not available")

SUPPORTED_EXTENSIONS = [".pdf", ".docx", ".txt", ".eml"]

class ParsedDocument:
    def __init__(self, chunks: List[str], meta: Optional[Dict] = None):
        self.chunks = chunks
        self.meta = meta or {}

def clean_text(text: str) -> str:
    """Clean and normalize extracted text"""
    if not text:
        return ""
    
    # Remove excessive whitespace
    text = re.sub(r'\s+', ' ', text)
    # Remove special characters that might cause issues
    text = re.sub(r'[^\w\s\.,;:!?()-]', ' ', text)
    # Remove excessive dots or dashes
    text = re.sub(r'\.{3,}', '...', text)
    text = re.sub(r'-{3,}', '---', text)
    
    return text.strip()

def chunk_text(text: str, max_chunk_size: int = 1024, overlap: int = 128) -> List[str]:
    """Chunk text into overlapping segments for semantic search."""
    if not text or not text.strip():
        return []
    
    # Clean the text first
    text = clean_text(text)
    
    if len(text) <= max_chunk_size:
        return [text] if text else []
    
    chunks = []
    start = 0
    
    while start < len(text):
        end = min(start + max_chunk_size, len(text))
        
        # Try to find a good breaking point (sentence boundary)
        if end < len(text):
            # Look for sentence endings within the last 100 characters
            search_start = max(start, end - 100)
            sentence_ends = []
            
            for pattern in ['. ', '! ', '? ', '; ', '\n\n']:
                pos = text.rfind(pattern, search_start, end)
                if pos != -1:
                    sentence_ends.append(pos + len(pattern))
            
            if sentence_ends:
                end = max(sentence_ends)
        
        chunk = text[start:end].strip()
        if chunk and len(chunk) > 10:  # Only add meaningful chunks
            chunks.append(chunk)
        
        start = max(start + 1, end - overlap)
    
    return chunks

def parse_pdf_pymupdf(path: str) -> List[str]:
    """Parse PDF using PyMuPDF (fitz) - most reliable option"""
    if not PYMUPDF_AVAILABLE:
        raise ImportError("PyMuPDF not available")
    
    try:
        doc = fitz.open(path)
        text_parts = []
        
        for page_num in range(len(doc)):
            page = doc.load_page(page_num)
            text = page.get_text()
            
            if text.strip():
                # Add page separator for context
                text_parts.append(f"[Page {page_num + 1}]\n{text}")
        
        doc.close()
        
        full_text = "\n\n".join(text_parts)
        return chunk_text(full_text)
        
    except Exception as e:
        raise Exception(f"PyMuPDF parsing failed: {str(e)}")

def parse_pdf_pypdf2(path: str) -> List[str]:
    """Fallback PDF parser using PyPDF2"""
    if not PYPDF2_AVAILABLE:
        raise ImportError("PyPDF2 not available")
    
    try:
        with open(path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            text_parts = []
            
            for page_num, page in enumerate(reader.pages):
                text = page.extract_text()
                if text.strip():
                    text_parts.append(f"[Page {page_num + 1}]\n{text}")
            
            full_text = "\n\n".join(text_parts)
            return chunk_text(full_text)
            
    except Exception as e:
        raise Exception(f"PyPDF2 parsing failed: {str(e)}")

def parse_docx(path: str) -> List[str]:
    """Parse DOCX files"""
    if not DOCX_AVAILABLE:
        raise ImportError("python-docx not available")
    
    try:
        doc = Document(path)
        paragraphs = []
        
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text.strip())
        
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join([cell.text.strip() for cell in row.cells])
                if row_text.strip():
                    paragraphs.append(row_text)
        
        full_text = "\n\n".join(paragraphs)
        return chunk_text(full_text)
        
    except Exception as e:
        raise Exception(f"DOCX parsing failed: {str(e)}")

def parse_text_file(path: str) -> List[str]:
    """Parse plain text files"""
    try:
        # Try UTF-8 first
        try:
            with open(path, 'r', encoding='utf-8') as f:
                text = f.read()
        except UnicodeDecodeError:
            # Fallback to latin-1
            with open(path, 'r', encoding='latin-1') as f:
                text = f.read()
        
        return chunk_text(text)
        
    except Exception as e:
        raise Exception(f"Text file parsing failed: {str(e)}")

def parse_email_simple(path: str) -> List[str]:
    """Simple email parsing without tika dependency"""
    try:
        with open(path, 'r', encoding='utf-8', errors='ignore') as f:
            content = f.read()
        
        # Basic email parsing - extract text content
        lines = content.split('\n')
        body_started = False
        body_lines = []
        
        for line in lines:
            # Skip headers until we find the body
            if not body_started:
                if line.strip() == '' or line.startswith('Content-Type:'):
                    body_started = True
                continue
            
            # Skip MIME boundaries and headers in body
            if line.startswith('--') or line.startswith('Content-'):
                continue
                
            body_lines.append(line)
        
        text = '\n'.join(body_lines)
        return chunk_text(text)
        
    except Exception as e:
        # Fallback to treating as plain text
        return parse_text_file(path)

def get_extension(path: str) -> str:
    """Get file extension from path, handling URLs"""
    parsed_url = urlparse(path)
    path_part = parsed_url.path
    return os.path.splitext(path_part)[-1].lower()

def parse_document(path: str) -> ParsedDocument:
    """Main document parsing function with comprehensive fallbacks"""
    ext = get_extension(path)
    
    try:
        if ext == ".pdf":
            # Try PyMuPDF first (most reliable)
            if PYMUPDF_AVAILABLE:
                try:
                    chunks = parse_pdf_pymupdf(path)
                    return ParsedDocument(chunks, {"parser": "pymupdf", "extension": ext})
                except Exception as e:
                    print(f"PyMuPDF failed: {e}")
            
            # Fallback to PyPDF2
            if PYPDF2_AVAILABLE:
                try:
                    chunks = parse_pdf_pypdf2(path)
                    return ParsedDocument(chunks, {"parser": "pypdf2", "extension": ext})
                except Exception as e:
                    print(f"PyPDF2 failed: {e}")
            
            raise Exception("No PDF parsers available")
            
        elif ext == ".docx":
            if DOCX_AVAILABLE:
                chunks = parse_docx(path)
                return ParsedDocument(chunks, {"parser": "docx", "extension": ext})
            else:
                raise Exception("python-docx not available")
            
        elif ext == ".eml":
            chunks = parse_email_simple(path)
            return ParsedDocument(chunks, {"parser": "email_simple", "extension": ext})
            
        elif ext in [".txt", ""]:
            chunks = parse_text_file(path)
            return ParsedDocument(chunks, {"parser": "text", "extension": ext})
            
        else:
            # Try to parse as text for unknown extensions
            try:
                chunks = parse_text_file(path)
                return ParsedDocument(chunks, {"parser": "text_fallback", "extension": ext})
            except Exception:
                raise ValueError(f"Unsupported file extension: {ext}")
                
    except Exception as e:
        print(f"Document parsing failed for {path}: {e}")
        # Return empty document with error info
        return ParsedDocument(
            chunks=[], 
            meta={
                "parser": "failed", 
                "error": str(e), 
                "extension": ext
            }
        )